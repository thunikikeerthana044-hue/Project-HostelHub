"""
Generate HostelHub Project Documentation (Word .docx)
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Helper styles ─────────────────────────────────────────────────────────────

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(30, 30, 30)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(26, 60, 94)  # HostelHub brand navy
    h.font.name = 'Calibri'

def add_table(headers, rows):
    """Add a styled table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()  # spacer

# ══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('HOSTELHUB')
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(26, 60, 94)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Online Hostel Booking & Management System')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline.add_run('A Full-Stack Web Application built with Django & Razorpay')
run.font.size = Pt(12)
run.font.italic = True
run.font.color.rgb = RGBColor(120, 120, 120)

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Project Documentation\nClassroom Presentation')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1.  Introduction',
    '2.  Problem Statement',
    '3.  Objectives',
    '4.  Technologies Used',
    '5.  System Architecture',
    '6.  Database Design (Models)',
    '7.  Module Descriptions',
    '    7.1  Accounts Module',
    '    7.2  Hostels Module',
    '    7.3  Booking & Payment Module',
    '    7.4  Review & Rating Module',
    '    7.5  Admin Panel Module',
    '8.  System Flow (User Journeys)',
    '    8.1  Student / Seeker Flow',
    '    8.2  Hostel Owner Flow',
    '9.  URL Routing Structure',
    '10. Key Features',
    '11. Project Directory Structure',
    '12. How to Run the Project',
    '13. Future Enhancements',
    '14. Conclusion',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading('1. Introduction', level=1)
doc.add_paragraph(
    'HostelHub is a full-stack web application designed to simplify the process of '
    'searching, comparing, and booking hostels, PGs (Paying Guests), and shared '
    'accommodations across cities in India. The platform serves two distinct user '
    'roles — Students/Travelers who are looking for accommodation, and Hostel Owners '
    'who want to list and manage their properties online.'
)
doc.add_paragraph(
    'The application provides a modern, responsive user interface with real-time '
    'search and filtering, secure online payment integration via Razorpay, '
    'role-based dashboards, a review & rating system, and a full admin panel for '
    'managing listings and bookings.'
)

# ══════════════════════════════════════════════════════════════════════════════
#  2. PROBLEM STATEMENT
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading('2. Problem Statement', level=1)
doc.add_paragraph(
    'Students and working professionals relocating to new cities often face '
    'significant challenges in finding reliable and affordable hostel accommodations. '
    'The current process typically involves:'
)
problems = [
    'Manually visiting multiple hostels, wasting time and effort.',
    'Lack of a centralized platform to compare prices, amenities, and reviews.',
    'No standardized booking process — most bookings happen through phone calls with no formal confirmation.',
    'Hostel owners have limited online visibility and depend on word-of-mouth referrals.',
    'No transparent review system to help new tenants make informed decisions.',
]
for p in problems:
    doc.add_paragraph(p, style='List Bullet')

doc.add_paragraph(
    'HostelHub addresses all of these pain points by providing a single unified platform '
    'that connects hostel seekers with property owners through a seamless digital experience.'
)

# ══════════════════════════════════════════════════════════════════════════════
#  3. OBJECTIVES
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading('3. Objectives', level=1)
objectives = [
    'Develop a web-based hostel booking platform with role-based access (Student & Owner).',
    'Implement advanced search and filtering by city, gender policy, price range, room type, and amenities.',
    'Integrate a secure online payment gateway (Razorpay) for token booking.',
    'Provide separate dashboards for students (booking history) and owners (property & booking management).',
    'Enable a review and 5-star rating system for verified hostels.',
    'Build a Django admin panel for super-admin oversight of all listings, bookings, and users.',
    'Design a responsive, modern UI using Bootstrap 5 and custom CSS.',
]
for o in objectives:
    doc.add_paragraph(o, style='List Bullet')

# ══════════════════════════════════════════════════════════════════════════════
#  4. TECHNOLOGIES USED
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading('4. Technologies Used', level=1)

add_table(
    ['Category', 'Technology', 'Purpose'],
    [
        ['Backend Framework', 'Django 6.0 (Python)', 'Server-side logic, ORM, URL routing, authentication'],
        ['Frontend', 'HTML5, CSS3, JavaScript', 'Page structure, styling, and interactive behavior'],
        ['UI Framework', 'Bootstrap 5', 'Responsive grid layout, components, and mobile-first design'],
        ['Icons', 'Bootstrap Icons (CDN)', 'Icon set used across the UI for visual clarity'],
        ['Typography', 'Google Fonts (Sora)', 'Modern, clean typeface for headings and branding'],
        ['Database', 'SQLite3', 'Lightweight relational database (default Django backend)'],
        ['Payment Gateway', 'Razorpay API', 'Secure online payment processing for bookings'],
        ['Image Handling', 'Pillow (PIL)', 'Image upload, validation, and processing for avatars & hostel photos'],
        ['Authentication', 'Django Auth System', 'Built-in user registration, login, session management'],
        ['Signals', 'Django Signals', 'Auto-creation of UserProfile when a new User is registered'],
        ['Template Engine', 'Django Template Language', 'Server-side HTML rendering with template inheritance'],
        ['Dev Server', 'Django runserver', 'Local development server with auto-reload'],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
#  5. SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading('5. System Architecture', level=1)
doc.add_paragraph(
    'HostelHub follows the Model-View-Template (MVT) architecture, which is Django\'s '
    'implementation of the classic MVC pattern:'
)

add_table(
    ['Layer', 'Component', 'Description'],
    [
        ['Model', 'models.py', 'Defines the database schema — UserProfile, Hostel, Room, Booking, Review, Amenity, HostelImage. Django ORM maps these classes to SQLite tables.'],
        ['View', 'views.py', 'Contains the business logic — handling HTTP requests, querying the database, processing form data, initiating payments, and returning responses.'],
        ['Template', 'templates/*.html', 'HTML files with Django Template Language that render the UI. Uses template inheritance (base.html → child templates).'],
        ['URL Dispatcher', 'urls.py', 'Maps URL patterns to view functions. The project has a root URL config that includes app-level URL configs.'],
        ['Forms', 'forms.py', 'Django form classes for user input validation — SignupForm, LoginForm, ProfileUpdateForm.'],
        ['Signals', 'signals.py', 'Django signals that auto-create a UserProfile whenever a new User object is saved.'],
        ['Admin', 'admin.py', 'Registers models with Django\'s built-in admin interface with custom list displays, filters, and search.'],
    ]
)

doc.add_paragraph('Request Lifecycle:', style='Heading 3')
doc.add_paragraph('1. User sends an HTTP request (e.g., GET /search/?city=Hyderabad).')
doc.add_paragraph('2. Django URL dispatcher matches the URL pattern to a view function.')
doc.add_paragraph('3. The view function queries the database via Django ORM.')
doc.add_paragraph('4. Data is passed to a Django template for HTML rendering.')
doc.add_paragraph('5. The rendered HTML response is sent back to the user\'s browser.')

# ══════════════════════════════════════════════════════════════════════════════
#  6. DATABASE DESIGN
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading('6. Database Design (Models)', level=1)
doc.add_paragraph(
    'The application uses 7 interconnected Django models. Below is the schema summary:'
)

doc.add_heading('6.1 UserProfile', level=2)
doc.add_paragraph('Extends Django\'s built-in User model with additional fields.')
add_table(
    ['Field', 'Type', 'Description'],
    [
        ['user', 'OneToOneField → User', 'Links to Django\'s auth User model'],
        ['role', 'CharField (choices)', 'Either "student" (Student/Traveler) or "owner" (Hostel Owner)'],
        ['phone', 'CharField', 'Contact phone number'],
        ['city', 'CharField', 'User\'s city'],
        ['avatar', 'ImageField', 'Profile picture upload'],
        ['created_at', 'DateTimeField', 'Account creation timestamp'],
    ]
)

doc.add_heading('6.2 Hostel', level=2)
doc.add_paragraph('Represents a hostel/PG listing with location, contact info, and amenities.')
add_table(
    ['Field', 'Type', 'Description'],
    [
        ['owner', 'ForeignKey → User', 'The owner who listed this hostel'],
        ['name', 'CharField', 'Hostel name'],
        ['description', 'TextField', 'Detailed description'],
        ['address, city, state, pincode', 'CharField/TextField', 'Location details'],
        ['phone, email', 'CharField/EmailField', 'Contact information'],
        ['gender_type', 'CharField (choices)', 'Male Only / Female Only / Mixed'],
        ['amenities', 'ManyToManyField → Amenity', 'WiFi, AC, Food, Laundry, etc.'],
        ['latitude, longitude', 'DecimalField', 'GPS coordinates (optional)'],
        ['is_verified', 'BooleanField', 'Admin verification status'],
        ['is_active', 'BooleanField', 'Whether the listing is active'],
        ['created_at, updated_at', 'DateTimeField', 'Timestamps'],
    ]
)

doc.add_heading('6.3 Room', level=2)
add_table(
    ['Field', 'Type', 'Description'],
    [
        ['hostel', 'ForeignKey → Hostel', 'Parent hostel'],
        ['room_type', 'CharField (choices)', 'Single / Double Sharing / Triple Sharing / Dormitory'],
        ['capacity', 'PositiveIntegerField', 'Number of beds'],
        ['price_per_month', 'DecimalField', 'Monthly rent'],
        ['deposit_amount', 'DecimalField', 'Security deposit'],
        ['total_rooms', 'PositiveIntegerField', 'Total rooms of this type'],
        ['available_rooms', 'PositiveIntegerField', 'Currently available count'],
        ['is_available', 'BooleanField', 'Availability flag'],
    ]
)

doc.add_heading('6.4 Booking', level=2)
add_table(
    ['Field', 'Type', 'Description'],
    [
        ['guest', 'ForeignKey → User', 'The student who made the booking'],
        ['room', 'ForeignKey → Room', 'The booked room'],
        ['check_in_date, check_out_date', 'DateField', 'Stay period'],
        ['status', 'CharField (choices)', 'Pending / Confirmed / Cancelled / Completed'],
        ['total_amount', 'DecimalField', 'Calculated total rent'],
        ['razorpay_order_id', 'CharField', 'Razorpay order reference'],
        ['razorpay_payment_id', 'CharField', 'Razorpay payment confirmation ID'],
        ['razorpay_signature', 'CharField', 'Razorpay signature for verification'],
        ['payment_status', 'CharField', 'unpaid / paid / failed'],
        ['special_requests', 'TextField', 'Optional guest notes'],
    ]
)

doc.add_heading('6.5 Review', level=2)
add_table(
    ['Field', 'Type', 'Description'],
    [
        ['hostel', 'ForeignKey → Hostel', 'The reviewed hostel'],
        ['reviewer', 'ForeignKey → User', 'The user who wrote the review'],
        ['rating', 'IntegerField (1–5)', '5-star rating with validators'],
        ['comment', 'TextField', 'Review text'],
        ['created_at', 'DateTimeField', 'Review timestamp'],
    ]
)
doc.add_paragraph('Constraint: unique_together = (hostel, reviewer) — one review per user per hostel.')

doc.add_heading('6.6 Amenity', level=2)
doc.add_paragraph('Reusable tags like WiFi, AC, Hot Water, Food, Laundry — linked to hostels via ManyToMany.')

doc.add_heading('6.7 HostelImage', level=2)
doc.add_paragraph('Multiple photos per hostel, with an is_main flag for the cover image.')

# ══════════════════════════════════════════════════════════════════════════════
#  7. MODULE DESCRIPTIONS
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading('7. Module Descriptions', level=1)

# 7.1
doc.add_heading('7.1 Accounts Module', level=2)
doc.add_paragraph('Handles all user authentication and profile management.')
add_table(
    ['Feature', 'View Function', 'URL', 'Description'],
    [
        ['User Registration', 'signup_view', '/signup/', 'Role-based signup form (Student or Owner). Uses Django\'s UserCreationForm with custom fields.'],
        ['User Login', 'login_view', '/login/', 'Username + password authentication. Redirects to appropriate dashboard.'],
        ['User Logout', 'logout_view', '/logout/', 'Ends the session and redirects to login page.'],
        ['Profile Management', 'profile_view', '/profile/', 'Edit first name, last name, email, phone, city, and avatar image.'],
    ]
)
doc.add_paragraph(
    'Key Implementation Detail: A Django Signal (post_save on User) automatically creates '
    'a UserProfile record whenever a new User is registered. This ensures every user always '
    'has an associated profile.'
)

# 7.2
doc.add_heading('7.2 Hostels Module — Search & Discovery', level=2)
doc.add_paragraph('The core module for browsing and finding hostels.')
add_table(
    ['Feature', 'View Function', 'URL', 'Description'],
    [
        ['Home Page', 'home', '/', 'Landing page with hero banner, search bar, popular cities quick-filter, featured hostels, stats counter, and "How It Works" section.'],
        ['Search & List', 'hostel_list', '/search/', 'Advanced search with filters: city, gender type, max price, room type, amenities, verified-only. Sorting by price, rating, or newest.'],
        ['Hostel Detail', 'hostel_detail', '/hostel/<id>/', 'Full detail page: photo gallery, description, amenities, room types with prices, review section, and booking form.'],
    ]
)
doc.add_paragraph(
    'Search Implementation: Uses Django ORM\'s Q objects for city/address text search, '
    'chained filter() calls for gender/amenity/verified filtering, Room subqueries for '
    'price/room-type filtering, and annotate() with Min/Avg for sorting by price and rating.'
)

# 7.3
doc.add_heading('7.3 Booking & Payment Module', level=2)
doc.add_paragraph('End-to-end booking flow with Razorpay payment integration.')
add_table(
    ['Step', 'View Function', 'Description'],
    [
        ['1. Initiate Payment', 'initiate_payment', 'Student selects room, check-in/out dates. System calculates total amount, creates a Razorpay Order, and renders the payment page.'],
        ['2. Payment Page', '(template)', 'Razorpay checkout widget is loaded. If API keys aren\'t configured, a mock/simulated payment flow is used for demo purposes.'],
        ['3. Verify Payment', 'verify_payment', 'AJAX POST endpoint. Verifies Razorpay signature (HMAC SHA256). On success: booking status → confirmed, payment → paid, room availability decremented.'],
        ['4. Success Page', 'booking_success', 'Shows booking confirmation with details — hostel name, room type, dates, amount paid.'],
        ['5. Failure Page', 'booking_failed', 'Shown if payment verification fails.'],
    ]
)
doc.add_paragraph(
    'Razorpay Integration: The system uses Razorpay\'s Orders API to create an order '
    'with the calculated amount in paise (INR × 100). The payment is captured automatically '
    '(payment_capture: 1). Signature verification uses HMAC-SHA256 with the Razorpay secret key. '
    'A fallback mock mode is available when API keys are not configured.'
)

# 7.4
doc.add_heading('7.4 Review & Rating Module', level=2)
doc.add_paragraph(
    'Students can submit a 1–5 star rating and a text comment for any hostel. '
    'The system enforces one review per user per hostel using the unique_together '
    'database constraint. If the user submits again, the existing review is updated '
    '(upsert pattern using get_or_create). The average rating is dynamically calculated '
    'and displayed on hostel cards and detail pages.'
)

# 7.5
doc.add_heading('7.5 Admin Panel Module', level=2)
doc.add_paragraph(
    'Django\'s built-in admin interface is configured at /admin/ with custom ModelAdmin '
    'classes for key models:'
)
add_table(
    ['Model', 'Admin Features'],
    [
        ['Hostel', 'List display: name, city, owner, verified, active. Filters: verified, active, city, gender. Inline editing of verified/active status.'],
        ['Room', 'List display: hostel, room type, price, availability.'],
        ['Booking', 'List display: guest, room, check-in, status, payment status, amount. Filters: status, payment status.'],
        ['Review', 'List display: hostel, reviewer, rating, date.'],
        ['UserProfile', 'Default admin interface.'],
        ['Amenity', 'Default admin interface.'],
        ['HostelImage', 'Default admin interface.'],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
#  8. SYSTEM FLOW
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading('8. System Flow (User Journeys)', level=1)

doc.add_heading('8.1 Student / Seeker Flow', level=2)
steps = [
    'Visit the home page → See featured hostels and popular city quick-links.',
    'Use the search bar → Filter by city, gender, price range, room type, amenities.',
    'Browse results → Sort by price, rating, or newest.',
    'Click on a hostel → View full details: photos, amenities, room options, reviews.',
    'Select a room → Choose check-in and check-out dates.',
    'Make payment → Razorpay checkout processes the token booking amount.',
    'Receive confirmation → Booking success page with all details.',
    'Visit Student Dashboard → View all past and current bookings with status.',
    'Write a review → Rate and comment on hostels you\'ve booked.',
    'Manage profile → Update personal info, phone, city, and avatar.',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('8.2 Hostel Owner Flow', level=2)
steps = [
    'Register as "Hostel Owner" → Signup with the owner role.',
    'Access Owner Dashboard → View all listed properties, earnings summary, active tenants.',
    'Add a New Hostel → Fill property details: name, address, city, gender policy, amenities, photo.',
    'Add Room Options → For each hostel, add room types (Single, Double, Triple, Dormitory) with pricing.',
    'Manage Bookings → View incoming bookings, confirm or cancel them.',
    'Track Earnings → Dashboard shows total earnings from confirmed bookings.',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

# ══════════════════════════════════════════════════════════════════════════════
#  9. URL ROUTING
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading('9. URL Routing Structure', level=1)
doc.add_paragraph('The project uses two app-level URL configurations included in the root urls.py:')

add_table(
    ['URL Pattern', 'View', 'Name', 'Auth Required'],
    [
        ['/', 'home', 'home', 'No'],
        ['/search/', 'hostel_list', 'hostel_list', 'No'],
        ['/hostel/<id>/', 'hostel_detail', 'hostel_detail', 'No'],
        ['/hostel/<id>/review/', 'submit_review', 'submit_review', 'Yes'],
        ['/book/<room_id>/', 'initiate_payment', 'initiate_payment', 'Yes'],
        ['/payment/verify/', 'verify_payment', 'verify_payment', 'Yes'],
        ['/booking/success/<id>/', 'booking_success', 'booking_success', 'Yes'],
        ['/booking/failed/', 'booking_failed', 'booking_failed', 'Yes'],
        ['/dashboard/', 'student_dashboard', 'student_dashboard', 'Yes'],
        ['/owner/dashboard/', 'owner_dashboard', 'owner_dashboard', 'Yes (Owner)'],
        ['/owner/hostel/add/', 'add_hostel', 'add_hostel', 'Yes (Owner)'],
        ['/owner/hostel/<id>/room/add/', 'add_room', 'add_room', 'Yes (Owner)'],
        ['/signup/', 'signup_view', 'signup', 'No'],
        ['/login/', 'login_view', 'login', 'No'],
        ['/logout/', 'logout_view', 'logout', 'No'],
        ['/profile/', 'profile_view', 'profile', 'Yes'],
        ['/admin/', 'Django Admin', 'admin', 'Superuser'],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
#  10. KEY FEATURES
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading('10. Key Features', level=1)

features = {
    'Role-Based Access Control': 'Two distinct user roles (Student and Owner) with separate dashboards, permissions, and UI experiences.',
    'Advanced Search & Filtering': 'Multi-criteria search combining city, gender policy, price range, room type, amenities, and verified-only filters with sorting options.',
    'Razorpay Payment Integration': 'Secure online payment with order creation, checkout widget, and cryptographic signature verification. Mock mode for demo environments.',
    'Responsive Design': 'Bootstrap 5-based mobile-first UI with custom CSS, glassmorphism effects, hover animations, and modern typography.',
    'Image Upload System': 'Hostel owners can upload property photos and users can set profile avatars. Handled via Django\'s ImageField and Pillow.',
    'Review & Rating System': 'One-review-per-user constraint, 5-star rating, average rating calculation, and chronological review display.',
    'Owner Property Management': 'Full CRUD for hostels and rooms: add listings, upload photos, manage room inventory, track bookings and earnings.',
    'Auto Profile Creation': 'Django signals automatically create a UserProfile when a new User registers, ensuring data integrity.',
    'Template Inheritance': 'A single base.html provides the navbar, footer, and meta structure. All pages extend this base for consistency.',
    'Django Admin Panel': 'Superuser access to manage all hostels, bookings, reviews, users, and amenities with custom list views and filters.',
}

for title, desc in features.items():
    doc.add_heading(title, level=3)
    doc.add_paragraph(desc)

# ══════════════════════════════════════════════════════════════════════════════
#  11. PROJECT DIRECTORY STRUCTURE
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading('11. Project Directory Structure', level=1)

tree = """hostelhub/                          ← Project Root
│
├── manage.py                       ← Django management script
├── db.sqlite3                      ← SQLite database file
├── run_project.bat                 ← One-click launcher script
├── seed_data.py                    ← Script to populate sample data
│
├── hostelhub/                      ← Django Project Configuration
│   ├── __init__.py
│   ├── settings.py                 ← All project settings (DB, apps, Razorpay keys)
│   ├── urls.py                     ← Root URL configuration
│   ├── wsgi.py                     ← WSGI entry point
│   └── asgi.py                     ← ASGI entry point
│
├── accounts/                       ← User Authentication App
│   ├── __init__.py
│   ├── apps.py                     ← App configuration
│   ├── forms.py                    ← SignupForm, LoginForm, ProfileUpdateForm
│   ├── models.py                   ← (UserProfile is in hostels/models.py)
│   ├── signals.py                  ← Auto-create UserProfile on User creation
│   ├── urls.py                     ← Auth URL patterns
│   ├── views.py                    ← signup, login, logout, profile views
│   └── templates/
│       ├── accounts/
│       │   ├── signup.html
│       │   ├── login.html
│       │   └── profile.html
│       └── base/
│           └── base.html           ← Master template (navbar, footer, CDN links)
│
├── hostels/                        ← Core Business Logic App
│   ├── __init__.py
│   ├── apps.py
│   ├── admin.py                    ← Admin panel configuration
│   ├── models.py                   ← All 7 data models
│   ├── urls.py                     ← Hostel & booking URL patterns
│   ├── views.py                    ← All view functions (16 views)
│   ├── tests.py                    ← Unit tests
│   ├── migrations/                 ← Database migration files
│   └── templates/hostels/
│       ├── home.html               ← Landing page
│       ├── hostel_list.html        ← Search results page
│       ├── hostel_detail.html      ← Individual hostel page
│       ├── payment.html            ← Razorpay checkout page
│       ├── booking_success.html    ← Confirmation page
│       ├── booking_failed.html     ← Payment failure page
│       ├── student_dashboard.html  ← Student bookings dashboard
│       ├── owner_dashboard.html    ← Owner management dashboard
│       ├── add_hostel.html         ← New hostel listing form
│       └── add_room.html           ← Add room to hostel form
│
└── media/                          ← User-uploaded files
    └── avatars/                    ← Profile pictures"""

p = doc.add_paragraph()
run = p.add_run(tree)
run.font.name = 'Consolas'
run.font.size = Pt(8.5)

# ══════════════════════════════════════════════════════════════════════════════
#  12. HOW TO RUN
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading('12. How to Run the Project', level=1)

doc.add_heading('Prerequisites', level=3)
prereqs = [
    'Python 3.10 or higher installed.',
    'pip package manager available.',
]
for p in prereqs:
    doc.add_paragraph(p, style='List Bullet')

doc.add_heading('Steps', level=3)
run_steps = [
    'Open a terminal/command prompt in the project root directory.',
    'Install dependencies:  pip install django razorpay pillow',
    'Run database migrations:  python manage.py migrate',
    'Start the development server:  python manage.py runserver',
    'Open browser at:  http://127.0.0.1:8000/',
    'OR simply double-click run_project.bat to automate all of the above.',
]
for i, s in enumerate(run_steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('Demo Login Credentials', level=3)
add_table(
    ['Role', 'Username', 'Password'],
    [
        ['Student / Seeker', 'seeker', 'password123'],
        ['Hostel Owner', 'owner', 'password123'],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
#  13. FUTURE ENHANCEMENTS
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading('13. Future Enhancements', level=1)
enhancements = [
    'Google Maps integration for hostel location display and distance-based search.',
    'Email/SMS notifications for booking confirmations and status updates.',
    'Real-time chat between students and hostel owners.',
    'Hostel comparison feature — side-by-side comparison of shortlisted hostels.',
    'Wishlist / Favorites — allow students to save hostels for later.',
    'Advanced analytics dashboard for owners — occupancy rates, revenue charts.',
    'Multi-language support for wider accessibility.',
    'Deployment to a production environment (AWS/Heroku) with PostgreSQL.',
    'Progressive Web App (PWA) support for mobile-like experience.',
]
for e in enhancements:
    doc.add_paragraph(e, style='List Bullet')

# ══════════════════════════════════════════════════════════════════════════════
#  14. CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading('14. Conclusion', level=1)
doc.add_paragraph(
    'HostelHub is a comprehensive, production-ready web application that demonstrates '
    'the effective use of Django\'s MVT architecture, relational database design, '
    'third-party API integration (Razorpay), and modern frontend technologies. '
    'The project successfully addresses the real-world problem of hostel discovery '
    'and booking by providing an intuitive, role-based platform for both students '
    'and property owners.'
)
doc.add_paragraph(
    'Key technical achievements include the implementation of a multi-criteria search '
    'engine with ORM-level query optimization, secure payment processing with '
    'cryptographic signature verification, a signal-driven user profile system, '
    'and a responsive UI built with Bootstrap 5 and custom CSS. The project is '
    'well-structured, maintainable, and designed for easy extension with future features.'
)

# ══════════════════════════════════════════════════════════════════════════════
#  SAVE
# ══════════════════════════════════════════════════════════════════════════════

output_path = os.path.join(os.path.expanduser('~'), 'Desktop', 'HostelHub_Project_Documentation.docx')
doc.save(output_path)
print(f"✅ Document saved successfully to: {output_path}")
